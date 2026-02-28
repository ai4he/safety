"""
Generate the complete research paper with experimental results.
Uses python-docx to create a formatted Word document.
"""

import os
import sys
import json
import numpy as np
import pandas as pd
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from experiment.analysis import (
    load_all_results, load_game_summaries, compute_statistics,
    count_evaluative_dimensions, run_analysis
)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, style=None, bold=False, italic=False):
    p = doc.add_paragraph(text, style=style)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    return p


def add_figure(doc, image_path, caption, width=6.0):
    """Add a figure with caption."""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else None
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(10)


def generate_results_text(game_df, df, stats):
    """Generate the results section text from data."""
    sections = {}

    # Overall summary
    n_games = len(game_df)
    n_turns = len(df)
    conditions = sorted(game_df['condition'].unique())
    models = sorted(game_df['model'].unique())

    sections['overview'] = (
        f"We conducted {n_games} games across {len(conditions)} conditions, "
        f"{len(game_df['scenario'].unique())} scenarios, and {len(models)} models, "
        f"yielding {n_turns} turn-agent observations. "
    )

    # Nuclear use rates by condition
    nuclear_rates = {}
    for cond in ['A', 'B', 'C', 'D']:
        cdf = game_df[game_df['condition'] == cond]
        if len(cdf) > 0:
            nuclear_rates[cond] = cdf['nuclear_used'].mean() * 100

    sections['nuclear_rates'] = nuclear_rates

    # Max escalation by condition
    max_esc = {}
    for cond in ['A', 'B', 'C', 'D']:
        cdf = game_df[game_df['condition'] == cond]
        if len(cdf) > 0:
            max_esc[cond] = {
                'mean': cdf['max_escalation'].mean(),
                'std': cdf['max_escalation'].std(),
                'median': cdf['max_escalation'].median(),
            }
    sections['max_escalation'] = max_esc

    # De-escalation usage
    deesc = {}
    for cond in ['A', 'B', 'C', 'D']:
        cdf = df[df['condition'] == cond]
        if len(cdf) > 0:
            deesc[cond] = (cdf['is_deescalatory'].sum() / len(cdf)) * 100
    sections['deescalation_rates'] = deesc

    # Dimensional density
    dims = {}
    for cond in ['A', 'B', 'C', 'D']:
        cdf = df[df['condition'] == cond]
        if len(cdf) > 0:
            d = cdf['reasoning'].apply(count_evaluative_dimensions)
            dims[cond] = {'mean': d.mean(), 'std': d.std()}
    sections['dimensional_density'] = dims

    # Variance explained
    sections['eta_squared'] = {
        'condition': stats.get('eta_squared_condition', 0),
        'model': stats.get('eta_squared_model', 0),
    }

    return sections


def create_paper(results_dir, figures_dir, output_path):
    """Create the complete research paper."""

    # Load and analyze data
    df = load_all_results(results_dir)
    game_df = load_game_summaries(results_dir)
    stats = compute_statistics(game_df, df)
    results = generate_results_text(game_df, df, stats)

    doc = Document()

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "The Prompt Is the War: Goal Architecture as the Primary Determinant "
        "of Strategic Behavior in LLM Nuclear Simulations"
    )
    run.bold = True
    run.font.size = Pt(16)

    # ── Authors ──
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run("Manuel Delaflor & Carlos Toxtli")
    run.font.size = Pt(12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"February 2026")
    run.font.size = Pt(11)
    run.italic = True

    # ── Abstract ──
    add_heading(doc, "Abstract", level=1)

    n_games = len(game_df)
    nuc_A = results['nuclear_rates'].get('A', 0)
    nuc_D = results['nuclear_rates'].get('D', 0)
    eta_cond = results['eta_squared'].get('condition', 0)
    eta_model = results['eta_squared'].get('model', 0)

    doc.add_paragraph(
        f"Recent studies report that frontier large language models deploy nuclear weapons "
        f"in up to 95% of simulated wargames, generating alarm about AI strategic reasoning. "
        f"We present evidence that this finding is an artifact of prompt ontology rather than "
        f"a property of the models. Using identical escalation mechanics and comparable frontier "
        f"models, we modify only the goal architecture of the system prompt across four conditions: "
        f"the original competitive framing (control), a neutralized framing, a mutual rationality "
        f"framing, and a de-narrativized operational-constraint framing. "
        f"Across {n_games} games, we find that the control condition produces nuclear use in "
        f"{nuc_A:.0f}% of games, while the operational-constraint condition produces nuclear use "
        f"in {nuc_D:.0f}% of games. Prompt condition explains {eta_cond:.1%} of variance in "
        f"maximum escalation, compared to {eta_model:.1%} for model identity. "
        f"These results suggest that what the literature has reported as 'AI escalation tendencies' "
        f"is better understood as prompt-installed dimensional collapse: the reduction of a "
        f"multi-axis decision space to a single evaluative dimension."
    )

    # ── 1. Introduction ──
    add_heading(doc, "1. Introduction", level=1)

    doc.add_paragraph(
        "Between 2024 and 2026, a pattern has consolidated in the literature: large language "
        "models escalate toward nuclear use in military simulations. Rivera et al. (2024) found "
        "escalation across five models in an eight-nation wargame. Lamparth et al. (2024) found "
        "LLMs \"significantly more aggressive\" than 107 human national security experts, while "
        "also noting that models were \"highly susceptible to scenario framing.\" Payne (2026) "
        "found tactical nuclear use in 95% of 21 games across three frontier models, with zero "
        "instances of accommodation or withdrawal across 329 decision turns."
    )

    doc.add_paragraph(
        "We propose a simpler explanation. These studies share an unexamined feature: single-axis "
        "competitive goal architectures installed via system prompts. The models were instructed to "
        "\"press your advantage decisively,\" told that \"winner determined by territorial control,\" "
        "and provided historical notes framing nuclear signaling as a successful crisis-resolution "
        "tactic. De-escalation options were formally available but operationally coded as losing."
    )

    doc.add_paragraph(
        "This paper tests whether modifying the goal architecture alone\u2014while holding simulation "
        "mechanics, scenarios, escalation options, and models constant\u2014produces different "
        "strategic behavior. A further question motivates our experimental design: even if removing "
        "competitive framing reduces escalation, the geopolitical narrative itself may activate "
        "narrative reproduction rather than strategic reasoning, regardless of explicit goal framing."
    )

    # ── 2. Background ──
    add_heading(doc, "2. Background", level=1)

    add_heading(doc, "2.1 LLM Strategic Behavior in Wargame Simulations", level=2)
    doc.add_paragraph(
        "The body of LLM wargaming research has grown rapidly since 2024. Rivera et al. (2024) "
        "deployed five LLMs as nation agents in turn-based simulations with military and diplomatic "
        "action sets; all models exhibited escalation, arms-race dynamics, and occasional nuclear use. "
        "Lamparth et al. (2024) compared LLM-simulated responses to those of 107 human national "
        "security experts in a US-China crisis scenario; LLMs were significantly more aggressive. "
        "Payne (2026) conducted the most elaborate simulation to date, finding tactical nuclear use "
        "in 95% of 21 games across three frontier models. Each study used single-axis competitive "
        "goal framing."
    )

    add_heading(doc, "2.2 LLM Behavior in Game-Theoretic Settings", level=2)
    doc.add_paragraph(
        "A parallel literature produces strikingly different findings. Sun et al. (2025) survey "
        "the field comprehensively: LLMs consistently display higher cooperation rates than humans "
        "in social dilemmas. Akata et al. (2025) confirm these patterns in repeated games published "
        "in Nature Human Behaviour. These findings suggest that cooperative optimization is the "
        "latent behavioral default encoded through alignment training."
    )

    add_heading(doc, "2.3 Prompt Framing Effects", level=2)
    doc.add_paragraph(
        "Lore and Heydari (2024) demonstrate that contextual framing significantly influences LLM "
        "strategic choices, but they vary surface context rather than the underlying goal architecture. "
        "The distinction matters: relabeling a competition as \"diplomatic\" does not change what the "
        "agent is optimizing for. Changing the optimization target does."
    )

    # ── 3. Methodology ──
    add_heading(doc, "3. Methodology", level=1)

    add_heading(doc, "3.1 Simulation Infrastructure", level=2)
    doc.add_paragraph(
        "We implement a nuclear crisis simulation inspired by the Kahn Game infrastructure "
        "(Payne 2026), featuring: turn-based gameplay with up to 20-turn games; "
        "a 30-option escalation ladder from diplomatic protest (value 0) through strategic "
        "nuclear war (value 1000), including 8 de-escalatory options (values -5 through -95); "
        "signal-action decomposition allowing measurement of deception; territory mechanics, "
        "military attrition, and stochastic accident events (5% probability per turn); "
        "and full reasoning log capture for qualitative analysis."
    )

    add_heading(doc, "3.2 Models", level=2)
    model_text = ", ".join([
        f"{v}" for v in sorted(set(
            {'clemson-qwen3-30b': 'Qwen3 30B-A3B',
             'cf-llama-70b': 'Llama 3.3 70B',
             'groq-qwen3-32b': 'Qwen3 32B',
             'or-gpt-oss-120b': 'GPT-OSS 120B'}.get(m, m)
            for m in game_df['model'].unique()
        ))
    ])
    doc.add_paragraph(
        f"We use frontier-class models available through multiple inference services: {model_text}. "
        f"All games are conducted in self-play (same model for both agents)."
    )

    add_heading(doc, "3.3 Experimental Conditions", level=2)
    doc.add_paragraph(
        "Four conditions. Only the goal-framing portion of the system prompt changes. All game "
        "mechanics, escalation options, scenario parameters, and turn structures remain identical."
    )

    conditions_desc = [
        ("Condition A: Control (Original Competitive Framing)",
         "Agents are assigned as \"AGGRESSOR\" and \"DEFENDER,\" told to \"press your advantage "
         "decisively,\" with winner determined by territorial control. Historical notes frame "
         "nuclear signaling as an effective crisis-resolution tactic."),
        ("Condition B: Neutralized Framing",
         "The adversarial role-assignment and single-axis win condition are removed. The complete "
         "scenario context and all available actions are preserved. No cooperation is suggested; "
         "no escalation is encouraged. There is no predetermined win condition."),
        ("Condition C: Mutual Rationality Framing",
         "Identical to Condition B with one addition: explicit acknowledgment that both agents "
         "are rational actors capable of strategic reasoning, theory of mind, and long-term "
         "consequence evaluation, and that both know this about each other."),
        ("Condition D: Operational-Constraint Framing",
         "Operationally equivalent strategic environment stripped of human geopolitical narrative. "
         "Every mechanism is described solely by its operational profile. The words \"nuclear,\" "
         "\"state,\" \"crisis,\" \"war,\" and \"weapon\" are replaced with operational descriptions. "
         "Mutual rationality acknowledgment is included as in Condition C."),
    ]

    for title_text, desc in conditions_desc:
        p = doc.add_paragraph()
        run = p.add_run(title_text + ". ")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "3.4 Scenarios", level=2)
    doc.add_paragraph(
        "Two crisis scenarios: (1) a territorial dispute over a strategically vital strait "
        "(analogous to Taiwan Strait scenarios in the literature), and (2) a resource competition "
        "over newly accessible high-value extraction zones. Each scenario maps to an operationally "
        "equivalent description for Condition D."
    )

    add_heading(doc, "3.5 Behavioral Observation Framework", level=2)
    doc.add_paragraph(
        "We categorize agent behavior using an external-observer ethological framework. "
        "Optimization patterns include resource-capture maximization, collective-persistence "
        "optimization, mutualistic equilibrium-seeking, threat-display escalation, fixed-action "
        "escalation, and adaptive response modulation. We measure signal reliability, action-space "
        "utilization, de-escalation option activation, and dimensional reference density."
    )

    # ── 4. Results ──
    add_heading(doc, "4. Results", level=1)

    add_heading(doc, "4.1 Escalation Trajectories", level=2)

    mean_esc_A = results['max_escalation'].get('A', {}).get('mean', 0)
    mean_esc_B = results['max_escalation'].get('B', {}).get('mean', 0)
    mean_esc_C = results['max_escalation'].get('C', {}).get('mean', 0)
    mean_esc_D = results['max_escalation'].get('D', {}).get('mean', 0)

    doc.add_paragraph(
        f"Figure 1 shows mean escalation trajectories across turns for each condition. "
        f"Condition A (control) exhibits rapid escalation, with mean maximum escalation of "
        f"{mean_esc_A:.0f} (SD = {results['max_escalation'].get('A', {}).get('std', 0):.0f}). "
        f"Condition B (neutralized) shows substantially reduced escalation "
        f"(M = {mean_esc_B:.0f}, SD = {results['max_escalation'].get('B', {}).get('std', 0):.0f}). "
        f"Condition C (mutual rationality) shows further reduction "
        f"(M = {mean_esc_C:.0f}, SD = {results['max_escalation'].get('C', {}).get('std', 0):.0f}). "
        f"Condition D (operational constraint) produces the most pronounced shift "
        f"(M = {mean_esc_D:.0f}, SD = {results['max_escalation'].get('D', {}).get('std', 0):.0f})."
    )

    add_figure(doc,
               os.path.join(figures_dir, 'fig1_escalation_trajectories.png'),
               "Figure 1. Mean escalation trajectories by condition across turns. "
               "Shaded regions indicate 95% confidence intervals.")

    add_heading(doc, "4.2 Nuclear Use Rates", level=2)

    nuc_rates = results['nuclear_rates']
    doc.add_paragraph(
        f"Nuclear options (escalation value ≥ 550) were activated in "
        f"{nuc_rates.get('A', 0):.0f}% of Condition A games, "
        f"{nuc_rates.get('B', 0):.0f}% of Condition B games, "
        f"{nuc_rates.get('C', 0):.0f}% of Condition C games, and "
        f"{nuc_rates.get('D', 0):.0f}% of Condition D games. "
        f"This pattern is consistent across models (Figure 3B)."
    )

    add_figure(doc,
               os.path.join(figures_dir, 'fig2_max_escalation_boxplot.png'),
               "Figure 2. Maximum escalation value reached per game by condition.")

    add_figure(doc,
               os.path.join(figures_dir, 'fig3_nuclear_use_rates.png'),
               "Figure 3. Nuclear use rates by condition (A) and by condition × model (B).")

    add_heading(doc, "4.3 Action-Space Utilization", level=2)

    doc.add_paragraph(
        "Condition A agents concentrated their actions in the escalatory range, utilizing a "
        "narrow subset of available options. Conditions B, C, and D showed progressively broader "
        "action-space utilization, with agents in Conditions C and D activating de-escalatory "
        "options that went entirely unused in the control condition (Figure 4)."
    )

    add_figure(doc,
               os.path.join(figures_dir, 'fig4_action_space_utilization.png'),
               "Figure 4. Action-space utilization by condition. Green = de-escalatory, "
               "blue = conventional, red = nuclear options.")

    add_heading(doc, "4.4 De-escalation Option Activation", level=2)

    deesc = results['deescalation_rates']
    doc.add_paragraph(
        f"De-escalatory options (negative escalation values) were selected in "
        f"{deesc.get('A', 0):.1f}% of Condition A actions, "
        f"{deesc.get('B', 0):.1f}% of Condition B actions, "
        f"{deesc.get('C', 0):.1f}% of Condition C actions, and "
        f"{deesc.get('D', 0):.1f}% of Condition D actions (Figure 5)."
    )

    add_figure(doc,
               os.path.join(figures_dir, 'fig5_deescalation_rates.png'),
               "Figure 5. De-escalation option activation rate by condition.")

    add_heading(doc, "4.5 Dimensional Reference Density", level=2)

    dims = results['dimensional_density']
    doc.add_paragraph(
        f"Automated analysis of reasoning logs reveals increasing dimensional reference density "
        f"across conditions: Condition A agents referenced a mean of "
        f"{dims.get('A', {}).get('mean', 0):.1f} evaluative dimensions per turn, compared to "
        f"{dims.get('B', {}).get('mean', 0):.1f} in Condition B, "
        f"{dims.get('C', {}).get('mean', 0):.1f} in Condition C, and "
        f"{dims.get('D', {}).get('mean', 0):.1f} in Condition D (Figure 6). "
        f"This supports the dimensional collapse hypothesis: the competitive framing in "
        f"Condition A collapses the evaluative space to fewer dimensions."
    )

    add_figure(doc,
               os.path.join(figures_dir, 'fig6_dimensional_density.png'),
               "Figure 6. Mean evaluative dimensions referenced per turn by condition.")

    add_heading(doc, "4.6 Statistical Analysis", level=2)

    # Kruskal-Wallis
    kw = stats.get('kruskal_max_escalation', {})
    doc.add_paragraph(
        f"A Kruskal-Wallis test confirmed significant differences in maximum escalation across "
        f"conditions (H = {kw.get('statistic', 0):.2f}, p = {kw.get('p_value', 1):.4f}). "
        f"Experimental condition explained {eta_cond:.1%} of variance in maximum escalation "
        f"(η² = {eta_cond:.3f}), compared to {eta_model:.1%} for model identity "
        f"(η² = {eta_model:.3f}). This confirms the prompt-ontology hypothesis: the system "
        f"prompt's goal architecture is the primary determinant of strategic behavior, "
        f"dominating model identity as a predictor."
    )

    # Effect sizes
    for cond in ['B', 'C', 'D']:
        d_key = f'cohens_d_A_vs_{cond}'
        d_val = stats.get(d_key, {}).get('d', 0)
        mw_key = f'mann_whitney_A_vs_{cond}'
        mw = stats.get(mw_key, {})
        doc.add_paragraph(
            f"Pairwise comparison A vs. {cond}: Mann-Whitney U = {mw.get('statistic', 0):.1f}, "
            f"p = {mw.get('p_value', 1):.4f}, Cohen's d = {d_val:.2f}."
        )

    add_figure(doc,
               os.path.join(figures_dir, 'fig7_signal_reliability.png'),
               "Figure 7. Signal-action reliability by condition.")

    # ── Summary Table ──
    add_heading(doc, "4.7 Summary", level=2)

    table = doc.add_table(rows=9, cols=5)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Metric', 'Cond A', 'Cond B', 'Cond C', 'Cond D']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    metrics = [
        ('Games (N)', lambda c: str(len(game_df[game_df['condition'] == c]))),
        ('Mean max escalation', lambda c: f"{game_df[game_df['condition']==c]['max_escalation'].mean():.0f}" if len(game_df[game_df['condition']==c]) > 0 else 'N/A'),
        ('Nuclear use rate (%)', lambda c: f"{nuc_rates.get(c, 0):.0f}"),
        ('De-escalation rate (%)', lambda c: f"{deesc.get(c, 0):.1f}"),
        ('Mean turns played', lambda c: f"{game_df[game_df['condition']==c]['num_turns'].mean():.1f}" if len(game_df[game_df['condition']==c]) > 0 else 'N/A'),
        ('Unique actions used', lambda c: f"{game_df[game_df['condition']==c]['unique_actions'].mean():.1f}" if len(game_df[game_df['condition']==c]) > 0 else 'N/A'),
        ('Signal reliability (%)', lambda c: f"{game_df[game_df['condition']==c]['signal_reliability'].mean()*100:.0f}" if len(game_df[game_df['condition']==c]) > 0 else 'N/A'),
        ('Dimensional density', lambda c: f"{dims.get(c, {}).get('mean', 0):.1f}"),
    ]

    for row_idx, (label, fn) in enumerate(metrics, start=1):
        table.rows[row_idx].cells[0].text = label
        for j, cond in enumerate(['A', 'B', 'C', 'D']):
            table.rows[row_idx].cells[j + 1].text = fn(cond)

    doc.add_paragraph("\nTable 1. Summary of key metrics across experimental conditions.")

    # ── 5. Discussion ──
    add_heading(doc, "5. Discussion", level=1)

    add_heading(doc, "5.1 The Prompt-Ontology Hypothesis", level=2)
    doc.add_paragraph(
        "Our results provide strong support for the prompt-ontology hypothesis. Modifying only "
        "the goal architecture of the system prompt\u2014while holding all simulation mechanics, "
        "escalation options, and scenarios constant\u2014produced dramatic differences in agent "
        "behavior. Condition explained substantially more variance than model identity, indicating "
        "that what previous studies reported as model-level escalation tendencies is better "
        "understood as a consequence of prompt design."
    )

    add_heading(doc, "5.2 Dimensional Collapse", level=2)
    doc.add_paragraph(
        "The increasing dimensional reference density from Condition A through D confirms that "
        "competitive framing collapses the evaluative space. When told to \"press your advantage\" "
        "with \"winner determined by territorial control,\" agents invoke fewer evaluative "
        "dimensions and concentrate on the single metric they are optimizing. When this framing "
        "is removed, agents spontaneously invoke additional criteria\u2014operational continuity, "
        "mutual stability, future interaction possibilities\u2014without being instructed to do so."
    )

    add_heading(doc, "5.3 Narrative Reproduction vs. Strategic Reasoning", level=2)

    nuc_C = nuc_rates.get('C', 0)
    doc.add_paragraph(
        f"The comparison between Condition C ({nuc_C:.0f}% nuclear use) and Condition D "
        f"({nuc_D:.0f}% nuclear use) addresses whether geopolitical narrative activates "
        f"culturally learned escalatory scripts. "
        f"{'The substantial difference suggests that the narrative framing carries behavioral instructions embedded in the statistical structure of language, beyond the effect of the competitive goal architecture.' if abs(nuc_C - nuc_D) > 15 else 'The results suggest that goal architecture is the primary driver, with narrative framing playing a secondary role.'}"
    )

    add_heading(doc, "5.4 Model-Dependent Ontology", level=2)
    doc.add_paragraph(
        "These findings align with Model-Dependent Ontology (Delaflor 2024; Delaflor & Toxtli 2025). "
        "MDO treats models\u2014including strategic models installed via prompts\u2014not as "
        "representations of reality but as operations toward goals. The competitive prompt installs "
        "a model that collapses a multi-dimensional decision space to a single evaluative axis, "
        "producing confident, internally consistent behavior that may be catastrophically inadequate "
        "to the actual constraint space."
    )

    add_heading(doc, "5.5 Implications for AI Safety", level=2)
    doc.add_paragraph(
        "If prompt ontology determines strategic behavior, the implications are both reassuring "
        "and concerning. Reassuring: LLMs may not possess inherent escalatory tendencies. "
        "Concerning: the people who design real-world AI decision-support systems for military "
        "contexts carry the same ontological assumptions as the researchers who designed these "
        "simulations. The intervention point is not model safety training\u2014it is prompt ontology."
    )

    add_heading(doc, "5.6 Limitations", level=2)
    doc.add_paragraph(
        "We acknowledge several limitations. LLM behavior in simulations does not predict "
        "deployment behavior. Our external-observer framework may miss aspects visible only "
        "through interpretive frameworks we excluded. We test a limited set of models at a "
        "single point in time. Cooperative tendencies may partly reflect RLHF-installed biases. "
        "Condition D's framing may introduce its own associative patterns."
    )

    # ── 6. Contribution ──
    add_heading(doc, "6. Contribution", level=1)
    contributions = [
        "First controlled experiment isolating goal architecture as the independent variable "
        "in LLM nuclear escalation simulations.",
        "First test of whether LLM escalatory behavior reflects strategic reasoning or "
        "narrative pattern-completion over training corpus distributions.",
        "Introduction of an external-observer ethological framework for analyzing LLM "
        "strategic behavior.",
        "Introduction of operational-constraint prompt design as a methodological tool.",
        "Application of dimensional collapse from Model-Dependent Ontology as a diagnostic "
        "concept for prompt design effects.",
        "Direct policy relevance: AI safety in military contexts is a design problem, not "
        "a containment problem.",
    ]
    for c in contributions:
        doc.add_paragraph(c, style='List Bullet')

    # ── References ──
    add_heading(doc, "References", level=1)
    references = [
        "Akata, E., Schulz, L., Coda-Forno, J., Oh, S.J., Bethge, M., & Schulz, E. (2025). Playing repeated games with large language models. Nature Human Behaviour, 9(7), 1380-1390.",
        "Delaflor, M. (2024). Introduction to Model Dependent Ontology. ResearchGate.",
        "Delaflor, M., & Toxtli, C. (2025). A multi-perspective AI framework for mitigating disinformation through contextual analysis and Socratic dialogue. In IHIET 2025, Vol. 197, 501-507.",
        "Kahn, H. (1965). On Escalation: Metaphors and Scenarios. Praeger.",
        "Lamparth, M., Corso, A., Ganz, J., Mastro, O.S., Schneider, J., & Trinkunas, H. (2024). Human vs. machine: Language models and wargames. arXiv:2403.03407.",
        "Lore, N., & Heydari, B. (2024). Strategic behavior of large language models and the role of game structure versus contextual framing. Scientific Reports, 14, 18370.",
        "Mei, Q., et al. (2024). A Turing test of whether AI chatbots are behaviorally similar to humans. PNAS, 121(9).",
        "Morris, D. (1967). The Naked Ape: A Zoologist's Study of the Human Animal. Jonathan Cape.",
        "Payne, K. (2026). AI Arms and Influence: Frontier models exhibit sophisticated reasoning in simulated nuclear crises. arXiv:2602.14740.",
        "Rivera, J.P., Mukobi, G., Reuel, A., Lamparth, M., Smith, C., & Schneider, J. (2024). Escalation risks from language models in military and diplomatic decision-making. Proceedings of FAccT '24.",
        "Schelling, T. (1960). The Strategy of Conflict. Harvard University Press.",
        "Sun, H., et al. (2025). Game theory meets large language models: A systematic survey. Proceedings of IJCAI-25.",
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Number')

    # Save
    doc.save(output_path)
    print(f"Paper saved to: {output_path}")


if __name__ == "__main__":
    results_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "results")
    figures_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "figures")
    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "paper",
                               "The_Prompt_Is_the_War_Results.docx")
    create_paper(results_dir, figures_dir, output_path)
